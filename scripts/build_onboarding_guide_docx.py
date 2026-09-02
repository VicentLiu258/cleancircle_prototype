from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs/onboarding-questionnaire-development-guide.md"
DOCX_PATH = ROOT / "docs/onboarding-questionnaire-development-guide.docx"

# Use one explicit Unicode font for both Latin and CJK runs. This avoids
# renderer-dependent fallback boxes when the document is opened in LibreOffice.
FONT = "Arial Unicode MS"
EAST_ASIA_FONT = "Arial Unicode MS"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5F6B7A"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
BORDER = "C8D0D9"


def set_run_font(run, size=10.5, color="1F2937", bold=None, italic=None, font=FONT):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_left_border(paragraph, color="2E74B5", size="18"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=MUTED)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("1F2937")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
        ("Heading 4", 11, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Code Block" not in styles:
        code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles["Code Block"]
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    code_style.font.size = Pt(8.5)
    code_style.font.color.rgb = RGBColor.from_string("243447")
    code_style.paragraph_format.space_before = Pt(2)
    code_style.paragraph_format.space_after = Pt(5)
    code_style.paragraph_format.line_spacing = 1.05


def add_rich_text(paragraph, text, size=10.5, color="1F2937"):
    # Keep inline emphasis readable without depending on Markdown-specific styles.
    token_re = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("``"):
            content = token[1:-1]
            run = paragraph.add_run(content)
            set_run_font(run, size=size - 0.4, color=INK, font="Consolas")
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, color=color, italic=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, color=color)


def parse_table(lines):
    rows = []
    for line in lines:
        stripped = line.strip().strip("|")
        parts = [part.strip() for part in stripped.split("|")]
        if all(re.fullmatch(r":?-{3,}:?", p.replace(" ", "")) for p in parts):
            continue
        rows.append(parts)
    max_cols = max(len(r) for r in rows)
    return [r + [""] * (max_cols - len(r)) for r in rows]


def table_widths(n):
    # Exact DXA widths; columns are intentionally uneven to protect narrative fields.
    presets = {
        2: [2700, 6660],
        3: [1900, 3500, 3960],
        4: [1500, 2900, 2500, 2460],
        5: [1350, 2450, 1900, 1900, 1760],
        6: [1250, 2050, 1700, 1600, 1450, 1310],
        7: [1150, 1850, 1600, 1350, 1250, 1100, 1060],
    }
    if n in presets:
        return presets[n]
    base = 9360 // n
    return [base] * n


def add_table(doc, rows):
    if not rows:
        return
    n = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n)
    set_table_geometry(table, table_widths(n))
    set_repeat_table_header(table.rows[0])
    for ri, row in enumerate(rows):
        for ci, value in enumerate(row):
            cell = table.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            if ri == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            add_rich_text(p, value, size=8.2 if n >= 5 else 8.7, color=INK if ri == 0 else "263342")
            if ri == 0:
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_callout(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.04)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.2
    set_paragraph_shading(p, LIGHT_GRAY)
    set_left_border(p)
    add_rich_text(p, text, size=10, color=INK)


def build_doc():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("NOT A PHASE · Onboarding / User Tagging Development Guide")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    r = footer.add_run("Internal development baseline · Page ")
    set_run_font(r, size=8.5, color=MUTED)
    add_page_field(footer)

    # First-page masthead, using the compact reference guide + memo masthead pattern.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("DEVELOPMENT GUIDE")
    set_run_font(r, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("NOT A PHASE Onboarding 问卷与用户标签系统")
    set_run_font(r, size=23, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("跨 API、Web 管理后台、移动端与推荐服务的统一开发基线")
    set_run_font(r, size=13, color=MUTED)

    for label, value in [
        ("版本", "V1.0（开发基线）"),
        ("日期", "2026-09-02"),
        ("依据", "NOT A PHASE Onboarding 问卷需求文档 V1（统一联动版）"),
        ("覆盖", "字段契约 · PostgreSQL · App API · Admin API · Web 管理后台 · 移动端 · 标签分组 · 推荐追踪"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        lr = p.add_run(f"{label}：")
        set_run_font(lr, size=10.5, color=INK, bold=True)
        vr = p.add_run(value)
        set_run_font(vr, size=10.5, color="1F2937")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(14)
    set_left_border(p, color=BLUE, size="12")
    add_rich_text(p, "使用方式：先冻结第 2 节的字段和枚举，再按第 4–8 节并行实现数据库、API、后台和移动端；源需求边界与本文技术实现的关系见第 0 节。", size=10.5, color=INK)

    raw = MD_PATH.read_text(encoding="utf-8")
    lines = raw.splitlines()
    i = 0
    content_started = False
    para_buffer = []
    table_buffer = []
    in_code = False
    code_buffer = []

    def flush_para():
        nonlocal para_buffer
        if para_buffer:
            text = " ".join(s.strip() for s in para_buffer).strip()
            if text:
                p = doc.add_paragraph()
                p.paragraph_format.keep_together = False
                add_rich_text(p, text)
            para_buffer = []

    def flush_table():
        nonlocal table_buffer
        if table_buffer:
            flush_para()
            add_table(doc, parse_table(table_buffer))
            table_buffer = []

    while i < len(lines):
        line = lines[i]
        # The Markdown source carries its own title/metadata for repository use;
        # the DOCX has a dedicated masthead above, so start at the first real
        # numbered section to avoid duplicating the cover block.
        if not content_started:
            if re.match(r"^##\s+0\.", line):
                content_started = True
            else:
                i += 1
                continue
        if line.startswith("```"):
            flush_para(); flush_table()
            if not in_code:
                in_code = True
                code_buffer = []
            else:
                p = doc.add_paragraph(style="Code Block")
                set_paragraph_shading(p, LIGHT_GRAY)
                for idx, code_line in enumerate(code_buffer):
                    if idx:
                        p.add_run().add_break()
                    r = p.add_run(code_line)
                    set_run_font(r, size=8.4, color="243447", font="Consolas")
                in_code = False
                code_buffer = []
            i += 1
            continue
        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        if line.startswith("|"):
            flush_para()
            table_buffer.append(line)
            i += 1
            continue
        if table_buffer and not line.startswith("|"):
            flush_table()

        heading = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading:
            flush_para()
            level = len(heading.group(1))
            text = heading.group(2).strip()
            # The Markdown title is already represented by the masthead.
            if level == 1 and text.startswith("NOT A PHASE"):
                i += 1
                continue
            actual_level = min(level, 4)
            p = doc.add_paragraph(style=f"Heading {actual_level}")
            add_rich_text(p, text, size={1: 16, 2: 13, 3: 12, 4: 11}[actual_level], color={1: BLUE, 2: BLUE, 3: DARK_BLUE, 4: DARK_BLUE}[actual_level])
            i += 1
            continue

        if line.startswith(">"):
            flush_para()
            add_callout(doc, line[1:].strip())
            i += 1
            continue

        bullet = re.match(r"^\s*[-*]\s+(.+)$", line)
        numbered = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if bullet or numbered:
            flush_para()
            p = doc.add_paragraph(style="List Bullet" if bullet else "List Number")
            add_rich_text(p, (bullet or numbered).group(1))
            i += 1
            continue

        if not line.strip():
            flush_para()
        else:
            para_buffer.append(line)
        i += 1

    flush_para(); flush_table()
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
    print(DOCX_PATH)
